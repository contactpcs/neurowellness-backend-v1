import os
import glob
import mne
import re
import threading

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use('Agg')  # headless backend — must be before any plt import
import matplotlib.pyplot as plt
import nibabel as nib


from datetime import datetime
from pathlib import Path

from mne import get_volume_labels_from_aseg
from mne.preprocessing import ICA
from mne.datasets import fetch_fsaverage
from mne.datasets import utils as mne_dataset_utils
from mne.viz import plot_topomap
from mne.transforms import apply_trans
from mne.minimum_norm import make_inverse_operator, apply_inverse_raw
from mne_connectivity import spectral_connectivity_epochs

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Image, Spacer,
    PageBreak, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from reportlab.lib.utils import simpleSplit
from reportlab.lib.units import cm


def _fetch_fsaverage_with_timeout(subjects_dir=None):
    original_downloader_params = mne_dataset_utils._downloader_params

    def _patched_downloader_params(*, auth=None, token=None):
        params = original_downloader_params(auth=auth, token=token)
        params["timeout"] = 300
        return params

    mne_dataset_utils._downloader_params = _patched_downloader_params
    try:
        return fetch_fsaverage(subjects_dir=subjects_dir, verbose=True)
    finally:
        mne_dataset_utils._downloader_params = original_downloader_params



import matplotlib.pyplot as plt
from matplotlib.offsetbox import OffsetImage, AnnotationBbox
from matplotlib.patches import FancyBboxPatch
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import subprocess as _subprocess


REPORT_DIR = " Report"
EEG_Bands  = {
    "Delta": (0.5, 4),
    "Theta": (4, 8),
    "Alpha1": (8, 10),
    "Alpha2": (10, 12),
    "Beta1": (12, 15),
    "Beta2": (15, 20),
    "Beta3": (20, 30),
    "Gamma": (30, 45)
}
# Current dir will have the all the nedf files of the patients, update it accordingly
CURRENT_DIR = Path(".")
TARGET_DIR = ''
OUTPUT_DOCX = "EEG_Report.docx"
OUTPUT_PDF = "EEG_Report.pdf"
OUTPUT_INDICATOR = r"EEG_indicators_loreta.pdf"
OUTPUT_BRAIN_CONNECTIVITY = "BrainFunctionConnectivity_Report.pdf"

_ASSETS_DIR = Path(__file__).parent / "assets"
ICONS_DIR = _ASSETS_DIR / "materials_and_icons"
brain_image = str(_ASSETS_DIR / "brain_stock_image.png")

all_dfs = {}
raw_mne_objs = {}
ch_names = []
dipoles = []

def load_brain_image(filename="brain.png"):
    return str(_ASSETS_DIR / filename)

normative_stats = {
    "integrated": {"mean": 0.18, "std": 0.05},
    "left_hemi": {"mean": 0.18, "std": 0.05},
    "right_hemi": {"mean": 0.19, "std": 0.05},
    "inter_hemi": {"mean": 0.16, "std": 0.04},
    "left_frontal": {"mean": 0.30, "std": 0.06},
    "right_frontal": {"mean": 0.32, "std": 0.06},
    "left_posterior": {"mean": 0.24, "std": 0.05},
    "right_posterior": {"mean": 0.25, "std": 0.05},
}

normative_abs_log = {
"Delta":  {"mean": -98.90, "std": 8.07},
"Theta":  {"mean": -106.85, "std": 6.35},
"Alpha1": {"mean": -110.78, "std": 7.02},
"Alpha2": {"mean": -106.58, "std": 5.65},
"Beta1":  {"mean": -113.08, "std": 5.96},
"Beta2":  {"mean": -112.93, "std": 6.30},
"Beta3":  {"mean": -111.30, "std": 6.10},
"Gamma":  {"mean": -111.78, "std": 7.39},
}

normative_rel = {
"Delta":  {"mean": 0.59, "std": 0.21},
"Theta":  {"mean": 0.09, "std": 0.03},

"Alpha1": {"mean": 0.06, "std": 0.06},
"Alpha2": {"mean": 0.12, "std": 0.09},

"Beta1":  {"mean": 0.02, "std": 0.02},
"Beta2":  {"mean": 0.03, "std": 0.02},
"Beta3":  {"mean": 0.04, "std": 0.04},

"Gamma":  {"mean": 0.04, "std": 0.06},
}

normative_abs_log_freq = {1: {'mean': -105.44157686572007, 'std': 7.3732786274124225},
    2: {'mean': -110.57402323218068, 'std': 7.391078018139551},
    3: {'mean': -113.68566294772434, 'std': 7.273577862602428},
    4: {'mean': -117.18990429989223, 'std': 6.671615937827277},
    5: {'mean': -117.21141142089412, 'std': 6.430242721594537},
    6: {'mean': -118.10824479471837, 'std': 4.994175125348294},
    7: {'mean': -117.81761271863353, 'std': 4.400066458082926},
    8: {'mean': -116.60337543313402, 'std': 6.241923921279898},
    9: {'mean': -117.67833166703879, 'std': 5.134707843147914},
    10: {'mean': -115.88690565567718, 'std': 6.4248203647620175},
    11: {'mean': -114.40437091507265, 'std': 4.9535739760443995},
    12: {'mean': -118.64672299229815, 'std': 4.192358289821445},
    13: {'mean': -119.0762472374891, 'std': 5.059602255849316},
    14: {'mean': -119.24017530912894, 'std': 4.1617312293947375},
    15: {'mean': -119.04821915230607, 'std': 5.213118400833486},
    16: {'mean': -119.52189857307805, 'std': 2.6186693630056084},
    17: {'mean': -119.23110158492985, 'std': 4.211430063638924},
    18: {'mean': -119.24442610314614, 'std': 4.1384486716894395},
    19: {'mean': -119.1991603030082, 'std': 4.3863796698801565},
    20: {'mean': -119.2363111602963, 'std': 4.182896044206641},
    21: {'mean': -119.42834450193519, 'std': 3.131086114119501},
    22: {'mean': -119.16525497805864, 'std': 4.5720867828242415},
    23: {'mean': -119.3739748647022, 'std': 3.428880881678219},
    24: {'mean': -119.16659209373366, 'std': 4.564763098652267},
    25: {'mean': -119.43810957044656, 'std': 3.077600631126901},
    26: {'mean': -119.33231143910379, 'std': 3.657080861910111},
    27: {'mean': -119.29540923178585, 'std': 3.859202575607868},
    28: {'mean': -119.24749581185279, 'std': 4.121635184653416},
    29: {'mean': -119.08217719067612, 'std': 5.027122564594458},
    30: {'mean': -119.42509736498411, 'std': 3.148871415673618},
    31: {'mean': -119.17394719869746, 'std': 4.5244775296373305},
    32: {'mean': -119.20552023870881, 'std': 4.351544867405041},
    33: {'mean': -119.13466919120933, 'std': 4.739612036788404},
    34: {'mean': -119.15336477275905, 'std': 4.637212119383762},
    35: {'mean': -119.2238087750201, 'std': 4.251374428590582},
    36: {'mean': -119.05212927823285, 'std': 5.1917017591056505},
    37: {'mean': -119.25146973939563, 'std': 4.0998690870823316},
    38: {'mean': -118.96558922490136, 'std': 5.6657011524792775},
    39: {'mean': -119.11172296514056, 'std': 4.865293693063243},
    40: {'mean': -119.11805115587143, 'std': 4.830632764948298},
    41: {'mean': -119.01209666458702, 'std': 5.410969414402788},
    42: {'mean': -119.2487919023765, 'std': 4.114536204489367},
    43: {'mean': -118.90430112142559, 'std': 6.001389920283158},
    44: {'mean': -118.92779469823175, 'std': 5.872710300551067},
    45: {'mean': -118.9841461167242, 'std': 5.56406086999367}}


normative_rel_freq = {1: {'mean': 0.16800225708331895, 'std': 0.07983249766340426},
    2: {'mean': 0.052334976747206545, 'std': 0.027574795703002426},
    3: {'mean': 0.02353466861816077, 'std': 0.011142109870495137},
    4: {'mean': 0.009442933716736697, 'std': 0.0039374420618586625},
    5: {'mean': 0.008780422014617291, 'std': 0.0023391590436634174},
    6: {'mean': 0.0071104455820899, 'std': 0.0018331510965745737},
    7: {'mean': 0.007775090095320198, 'std': 0.006104209392634908},
    8: {'mean': 0.011555500255525637, 'std': 0.010038528024213489},
    9: {'mean': 0.00792929501693211, 'std': 0.006949670922846149},
    10: {'mean': 0.02062017459454562, 'std': 0.024984913048727537},
    11: {'mean': 0.026666372126656058, 'std': 0.02222056560661193},
    12: {'mean': 0.006574971578523256, 'std': 0.005023153178167516},
    13: {'mean': 0.002886003269541601, 'std': 0.0013900055113263612},
    14: {'mean': 0.0022621806353929335, 'std': 0.0013522941365074564},
    15: {'mean': 0.002217510680933273, 'std': 0.0015180071796920592},
    16: {'mean': 0.0017977587181490144, 'std': 0.0013870318710485342},
    17: {'mean': 0.0019023288273921296, 'std': 0.001338900260227027},
    18: {'mean': 0.0016367073451233544, 'std': 0.001093006100404076},
    19: {'mean': 0.001578065947796072, 'std': 0.0010420641456668305},
    20: {'mean': 0.0016219590701626314, 'std': 0.0010820238646351138},
    21: {'mean': 0.001572956680771569, 'std': 0.00109233100380205},
    22: {'mean': 0.0017352409772921587, 'std': 0.0007272626797929751},
    23: {'mean': 0.0013138680027141162, 'std': 0.0006390974093394778},
    24: {'mean': 0.0009685814794306037, 'std': 0.0005404046657466789},
    25: {'mean': 0.0009907418108825537, 'std': 0.0006981792349889366},
    26: {'mean': 0.0008720206654913852, 'std': 0.0004949701991187486},
    27: {'mean': 0.0010734012908072122, 'std': 0.0020252898572440028},
    28: {'mean': 0.0006937875836987787, 'std': 0.00036886696839508507},
    29: {'mean': 0.0007274384797575442, 'std': 0.0004279225475524508},
    30: {'mean': 0.0006888880983827988, 'std': 0.00046932042759942403},
    31: {'mean': 0.0009244483288483669, 'std': 0.0015699877862591794},
    32: {'mean': 0.0006732617959423376, 'std': 0.00041159324169291376},
    33: {'mean': 0.0006170228070254625, 'std': 0.0003232770787123712},
    34: {'mean': 0.0007026243273251619, 'std': 0.0007057972466421216},
    35: {'mean': 0.0006222316749009075, 'std': 0.00043774152938040313},
    36: {'mean': 0.0005784454437099576, 'std': 0.0004308718966187681},
    37: {'mean': 0.0008612299740563409, 'std': 0.001848351094279595},
    38: {'mean': 0.0007351281260139276, 'std': 0.0007508784863392669},
    39: {'mean': 0.0005973827902575645, 'std': 0.00038256823248154976},
    40: {'mean': 0.0012456225387530104, 'std': 0.0037613272553123443},
    41: {'mean': 0.0006977645100421062, 'std': 0.0005957144751831503},
    42: {'mean': 0.0009574017945832261, 'std': 0.0021471573904290695},
    43: {'mean': 0.0016414035349712642, 'std': 0.0050261380396187764},
    44: {'mean': 0.001295875413928887, 'std': 0.002105539221939472},
    45: {'mean': 0.006824309569524927, 'std': 0.03098086304073419}}


# Expensive MNE resources — initialized lazily on first analysis job, NOT at import time.
fsaverage_path = None
trans = None
subjects_dir = None
subject = "fsaverage"
montage = None
src = None
bem = None
bem_sol = None

_init_lock = threading.Lock()
_resources_ready = False


def _init_resources() -> None:
    """Load fsaverage + compute BEM/source-space once. Thread-safe."""
    global fsaverage_path, trans, subjects_dir, montage, src, bem, bem_sol, _resources_ready
    if _resources_ready:
        return
    with _init_lock:
        if _resources_ready:
            return
        import mne as _mne
        fsaverage_path = _fetch_fsaverage_with_timeout()
        trans = os.path.join(fsaverage_path, "bem", "fsaverage-trans.fif")
        subjects_dir = os.path.dirname(fsaverage_path)
        montage = _mne.channels.make_standard_montage("standard_1020")
        src = _mne.setup_source_space(
            subject=subject,
            spacing="oct6",
            subjects_dir=subjects_dir,
            add_dist=False,
        )
        bem = _mne.make_bem_model(
            subject=subject,
            ico=4,
            conductivity=(0.3, 0.006, 0.3),
            subjects_dir=subjects_dir,
        )
        bem_sol = _mne.make_bem_solution(bem)
        _resources_ready = True

# HELPER FUNCTIONS
def doc_to_pdf(input_doc, output_pdf, target_dir):
    input_path = f"{target_dir}/{input_doc}"
    result = _subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", target_dir, input_path],
        capture_output=True, text=True, timeout=120,
    )
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")


def preprocess_raw(raw):
    raw = raw.copy()
    raw.load_data()

    if 'EXT' in raw.ch_names:
        raw.set_channel_types({'EXT': 'misc'})
    raw.pick_types(eeg=True, exclude='bads')
    raw.set_eeg_reference('average', projection=False)
    raw.filter(1., 45., fir_design='firwin')
    return raw


def make_epochs(raw):
    epochs = mne.make_fixed_length_epochs(
        raw,
        duration=2.0,
        overlap=1.0,
        preload=True
    )

    reject_criteria = dict(eeg=150e-6)
    epochs.drop_bad(reject=reject_criteria)

    return epochs


def define_regions(ch_names):
    idx = {ch: i for i, ch in enumerate(ch_names)}

    left_hemi = ['Fp1','F7','F3','C3','P7','P3','O1','T7']
    right_hemi = ['Fp2','F8','F4','C4','P8','P4','O2','T8']

    left_frontal = ['Fp1','F7','F3']
    right_frontal = ['Fp2','F8','F4']

    left_posterior = ['P7','P3','O1']
    right_posterior = ['P8','P4','O2']

    return {
        "left_hemi": [idx[ch] for ch in left_hemi if ch in idx],
        "right_hemi": [idx[ch] for ch in right_hemi if ch in idx],
        "inter_hemi": (
            [idx[ch] for ch in left_hemi if ch in idx],
            [idx[ch] for ch in right_hemi if ch in idx],
        ),
        "left_frontal": [idx[ch] for ch in left_frontal if ch in idx],
        "right_frontal": [idx[ch] for ch in right_frontal if ch in idx],
        "left_posterior": [idx[ch] for ch in left_posterior if ch in idx],
        "right_posterior": [idx[ch] for ch in right_posterior if ch in idx],
    }



def coherence_band(epochs, fmin, fmax):
    csd = mne.time_frequency.csd_multitaper(
        epochs,
        fmin=fmin,
        fmax=fmax,
        adaptive=True,
        normalization='full',
        verbose=False
    )

    # Extract cross-spectra
    csd_data = csd.get_data()  # (n_ch, n_ch)

    # Auto spectra
    psd = np.diag(csd_data)

    # Coherence formula
    coh = np.abs(csd_data) ** 2 / (
        psd[:, None] * psd[None, :]
    )

    np.fill_diagonal(coh, 0.0)
    return coh

def compute_band_connectivity(epochs, fmin, fmax):
    con = spectral_connectivity_epochs(
        epochs,
        method="coh",
        sfreq=epochs.info["sfreq"],
        fmin=fmin,
        fmax=fmax,
        faverage=True,
        verbose=False
    )

    n_ch = len(epochs.ch_names)
    data = con.get_data()

    # Handle flattened square case (e.g. 361 × 1)
    if data.shape == (n_ch * n_ch, 1):
        mat = data.reshape((n_ch, n_ch), order = "F")
    elif data.shape == (n_ch, n_ch):
        mat = data
    else:
        raise ValueError(f"Unexpected connectivity shape: {data.shape}")
    mat = 0.5 * (mat + mat.T)
    # Remove diagonal (critical for coherence)
    np.fill_diagonal(mat, np.nan)

    return mat

def t_score(value, mean, std):
    return 50 + 10 * (value - mean) / std

def classify(t):
    if t < 40:
        return "Low"
    elif t > 60:
        return "High"
    return "Normal"

def analytics_text(region, t_score, label):
    base = f"{region.replace('_',' ').title()} brain functional connectivity is {label.lower()} ({int(t_score)})."

    MAP = {
        "left_frontal": "Higher-order cognitive functioning, executive functioning, planning, decision-making, and behavioral regulation are",
        "right_frontal": "Emotion regulation, social cognition, intuitive thinking, and processing of nonverbal information is",
        "left_posterior": "Visual perception, language processing, auditory information processing, and memory formation are functioning at",
        "right_posterior": "Spatial awareness, visual information processing, nonverbal memory, and sensory information integration is"
    }

    if region in MAP:
        return base + "\n" + MAP[region] + f" {label.lower()} levels."

    return base

def tscore_bar(ax, t_score, title=None):
    ax.set_facecolor("white")

    bar_y = 0.5
    bar_h = 0.15

    # Main rounded bar
    bar = FancyBboxPatch(
        (0, bar_y - bar_h / 2),
        100,
        bar_h,
        boxstyle="round,pad=0.02,rounding_size=0.30",
        linewidth=0,
        facecolor="#e0e0e0",
        zorder=1
    )
    ax.add_patch(bar)

    # Red breakpoints
    for x in (40, 60):
        ax.plot([x, x], [bar_y - bar_h / 2, bar_y + bar_h / 2],
                color="red", linewidth=2, zorder=2)

    # Triangle pointer
    ax.plot(t_score, bar_y + bar_h / 2 + 0.05,
            marker="v", color="black", markersize=8, zorder=3)

    # Numeric label
    ax.text(t_score, bar_y + bar_h / 2 + 0.15,
            f"{int(round(t_score))}",
            ha="center", va="bottom", fontsize=10, fontweight="bold")

    # Labels
    ax.text(20, bar_y - 0.25, "Low", ha="center", fontsize=8)
    ax.text(50, bar_y - 0.25, "Normal", ha="center", fontsize=8)
    ax.text(80, bar_y - 0.25, "High", ha="center", fontsize=8)

    if title:
        ax.set_title(title, fontsize=10, pad=4)

    ax.set_xlim(0, 100)
    ax.set_ylim(0, 1)
    ax.axis("off")

def tscore_bar_vertical(ax, t_score, title=None):
    ax.set_facecolor("white")

    # Geometry
    bar_x = 0.5
    bar_w = 0.18

    # -----------------------------
    # Main vertical rounded bar
    # -----------------------------
    bar = FancyBboxPatch(
        (bar_x - bar_w / 2, 0),
        bar_w,
        100,
        boxstyle="round,pad=0.02,rounding_size=0.12",
        linewidth=0,
        facecolor="#e6e6e6",
        zorder=1
    )
    ax.add_patch(bar)

    # -----------------------------
    # Red breakpoints (40, 60)
    # -----------------------------
    for y in (40, 60):
        ax.plot(
            [bar_x - bar_w / 2, bar_x + bar_w / 2],
            [y, y],
            color="red",
            linewidth=2,
            zorder=2
        )

    # -----------------------------
    # Pointer (triangle → into bar)
    # -----------------------------
    ax.plot(
        bar_x + bar_w / 2 + 0.08,
        t_score,
        marker="<",
        color="black",
        markersize=9,
        zorder=3
    )

    # -----------------------------
    # Numeric value (top)
    # -----------------------------
    ax.text(
        bar_x,
        104,
        f"{int(round(t_score))}",
        ha="center",
        va="bottom",
        fontsize=12,
        fontweight="bold",
        color="red" if t_score < 40 or t_score > 60 else "black"
    )

    # -----------------------------
    # Labels inside bar
    # -----------------------------
    ax.text(bar_x, 80, "High", ha="center", va="center", fontsize=9)
    ax.text(bar_x, 50, "Normal", ha="center", va="center", fontsize=9)
    ax.text(bar_x, 20, "Low", ha="center", va="center", fontsize=9)

    # -----------------------------
    # Title
    # -----------------------------
    if title:
        ax.set_title(title, fontsize=10, pad=6)

    # -----------------------------
    # Axis cleanup
    # -----------------------------
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 110)
    ax.axis("off")


def mean_connectivity(mat, idx_a, idx_b=None):
    values = []

    if idx_b is None:
        for i in idx_a:
            for j in idx_a:
                if i < j:
                    values.append(mat[i, j])
    else:
        for i in idx_a:
            for j in idx_b:
                values.append(mat[i, j])

    values = np.array(values, dtype=float)
    values = values[~np.isnan(values)]

    if len(values) == 0:
        return np.nan

    return float(values.mean())

def draw_header(canvas, doc, metadata):
    styles = getSampleStyleSheet()

    header_text = f"""
    <b>Client name</b>: {metadata['name']} &nbsp;&nbsp;&nbsp;
    <b>Sex</b>: {metadata['sex']}<br/>
    <b>Date of Birth (age)</b>: {metadata['date_of_birth']} ({metadata['age']}Y) &nbsp;&nbsp;&nbsp;
    <b>EEG acq. date</b>: {metadata['meas_date']}
    """

    p = Paragraph(header_text, styles["Normal"])
    w, h = p.wrap(doc.width, doc.topMargin)
    p.drawOn(canvas, doc.leftMargin, doc.height + doc.topMargin - h + 10)



def add_intro_text(flow):
    styles = getSampleStyleSheet()
    normal = styles["Normal"]

    intro_paragraphs = [
        """Functional connectivity of the brain is closely related to brain plasticity and resilience.
        Brain plasticity refers to the brain’s ability to change structurally and functionally in
        response to environmental changes, learning, experience, or injury. This process of neural
        network modulation and reorganization can be described through functional connectivity
        measured using quantitative EEG analysis.""",

        """In general, a brain that has undergone healthy development maintains balance within and
        between functional regions. However, the brain may reinforce connectivity in specific
        regions when adapting to environmental demands or when structural or functional
        modifications are required. This adaptive capability is an inherent characteristic of the
        human brain.""",

        """Individuals may show variability in functional connectivity patterns depending on
        long-term learning, professional training, and environmental exposure. For example,
        individuals engaged in analytical or logical tasks may demonstrate relatively stronger
        connectivity in left-hemispheric networks, whereas individuals engaged in artistic or
        social activities may show relatively stronger right-hemispheric connectivity.""",

        """Excessively increased connectivity may be associated with inefficient information
        processing, excessive mental effort, or reduced flexibility of neural networks.
        Conversely, insufficient connectivity may be associated with reduced cognitive efficiency
        and difficulties in coordination between brain regions.""",

        """Brain functional connectivity analysis in this report is presented in three categories:
        integrated functional connectivity of the whole brain, hemispheric connectivity within and
        between the left and right hemispheres, and local connectivity within frontal and posterior
        brain regions. These indices provide information related to network balance and functional
        organization of the brain."""
    ]

    for p in intro_paragraphs:
        flow.append(Paragraph(p, normal))
        flow.append(Spacer(1, 12))

def integrated_analytics_text(t, c):
    return f"""
    Both the left hemisphere ({int(t['left_hemi'])}), and right hemisphere ({int(t['right_hemi'])})
    have {c['left_hemi'].lower()} connectivity.<br/><br/>

    • The connectivity between the right hemisphere and left hemisphere is
    {c['inter_hemi'].lower()} ({int(t['inter_hemi'])}).<br/>
    • Left frontal brain functional connectivity is
    {c['left_frontal'].lower()} ({int(t['left_frontal'])}).<br/>
    • Right frontal brain functional connectivity is
    {c['right_frontal'].lower()} ({int(t['right_frontal'])}).<br/>
    • Left posterior brain functional connectivity is
    {c['left_posterior'].lower()} ({int(t['left_posterior'])}).<br/>
    • Right posterior brain functional connectivity is
    {c['right_posterior'].lower()} ({int(t['right_posterior'])}).
    """


def hemispheric_analytics_text(t, c):
    return f"""
    Both the left hemisphere ({int(t['left_hemi'])}), and right hemisphere ({int(t['right_hemi'])})
    have {c['left_hemi'].lower()} connectivity.<br/><br/>

    The connectivity between the right hemisphere and left hemisphere is
    {c['inter_hemi'].lower()} ({int(t['inter_hemi'])}).
    """


def local_analytics_text(t, c):
    return f"""
    • Left frontal brain functional connectivity is
    {c['left_frontal'].lower()} ({int(t['left_frontal'])}).<br/>
    • Right frontal brain functional connectivity is
    {c['right_frontal'].lower()} ({int(t['right_frontal'])}).<br/>
    • Left posterior brain functional connectivity is
    {c['left_posterior'].lower()} ({int(t['left_posterior'])}).<br/>
    • Right posterior brain functional connectivity is
    {c['right_posterior'].lower()} ({int(t['right_posterior'])}).
    """


def build_pdf(output_path, sections, metadata):
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=72,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()
    flow = []

    # ---------- TITLE ----------
    flow.append(Paragraph(
        "<b>Brain Function Connectivity</b>",
        styles["Title"]
    ))
    flow.append(Spacer(1, 16))

    # ---------- INTRO ----------
    add_intro_text(flow)

    # =========================
    # SECTION 1: INTEGRATED
    # =========================
    flow.append(PageBreak())
    flow.append(Paragraph(
        "1 Integrated Brain Functional Connectivity",
        styles["Heading2"]
    ))
    flow.append(Spacer(1, 12))

    flow.append(Image(
        sections["integrated"]["figure"],
        width=400,
        height=120
    ))
    flow.append(Spacer(1, 12))

    flow.append(Paragraph(
        "<b>Analytics results</b>",
        styles["Normal"]
    ))
    flow.append(Spacer(1, 6))

    flow.append(Paragraph(
        integrated_analytics_text(
            sections["integrated"]["t_scores"],
            sections["integrated"]["classes"]
        ),
        styles["Normal"]
    ))

    # =========================
    # SECTION 2: HEMISPHERIC
    # =========================
    flow.append(PageBreak())
    flow.append(Paragraph(
        "2 Hemispheric Connectivity",
        styles["Heading2"]
    ))
    flow.append(Spacer(1, 12))

    flow.append(Image(
        sections["hemispheric"]["figure"],
        width=420,
        height=260
    ))
    flow.append(Spacer(1, 12))

    flow.append(Paragraph(
        "<b>Analytics results</b>",
        styles["Normal"]
    ))
    flow.append(Spacer(1, 6))

    flow.append(Paragraph(
        hemispheric_analytics_text(
            sections["hemispheric"]["t_scores"],
            sections["hemispheric"]["classes"]
        ),
        styles["Normal"]
    ))

    # =========================
    # SECTION 3: LOCAL
    # =========================
    flow.append(PageBreak())
    flow.append(Paragraph(
        "3 Local Connectivity",
        styles["Heading2"]
    ))
    flow.append(Spacer(1, 12))

    flow.append(Image(
        sections["local"]["figure"],
        width=420,
        height=260
    ))
    flow.append(Spacer(1, 12))

    flow.append(Paragraph(
        "<b>Analytics results</b>",
        styles["Normal"]
    ))
    flow.append(Spacer(1, 6))

    flow.append(Paragraph(
        local_analytics_text(
            sections["local"]["t_scores"],
            sections["local"]["classes"]
        ),
        styles["Normal"]
    ))

    # ---------- DISCLAIMER ----------
    flow.append(PageBreak())
    flow.append(Spacer(1, 24))
    flow.append(Paragraph(
        "This result is based on AI-assisted quantitative EEG analysis using default normative "
        "values and may differ from actual clinical conditions. This report is not a diagnosis.",
        styles["Italic"]
    ))

    doc.build(
        flow,
        onFirstPage=lambda c, d: draw_header(c, d, metadata),
        onLaterPages=lambda c, d: draw_header(c, d, metadata),
    )

def plot_hemispheric_composite(
    t_left,
    t_right,
    t_inter,
    brain_img_path,
    save_path
):
    import matplotlib.pyplot as plt
    import matplotlib.image as mpimg

    fig = plt.figure(figsize=(8, 5))
    fig.patch.set_facecolor("white")

    # -----------------------------
    # Left hemisphere indicator
    # -----------------------------
    ax_left = fig.add_axes([0.02, 0.45, 0.30, 0.40])
    tscore_bar_vertical(ax_left, t_left, "Left hemisphere")

    # -----------------------------
    # Brain image (center)
    # -----------------------------
    ax_brain = fig.add_axes([0.36, 0.42, 0.28, 0.46])
    ax_brain.imshow(mpimg.imread(brain_img_path))
    ax_brain.axis("off")

    # -----------------------------
    # Right hemisphere indicator
    # -----------------------------
    ax_right = fig.add_axes([0.68, 0.45, 0.30, 0.40])
    tscore_bar_vertical(ax_right, t_right, "Right hemisphere")

    # -----------------------------
    # Inter-hemispheric indicator (below brain)
    # -----------------------------
    ax_inter = fig.add_axes([0.20, 0.12, 0.60, 0.22])
    tscore_bar(ax_inter, t_inter, "Inter hemispheric")

    # -----------------------------
    # Save
    # -----------------------------
    plt.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close()

def plot_local_composite(
    t_lf, t_rf, t_lp, t_rp,
    brain_img_path, save_path
):
    import matplotlib.pyplot as plt
    import matplotlib.image as mpimg

    fig = plt.figure(figsize=(8, 8))
    fig.patch.set_facecolor("white")

    # Left frontal
    ax_lf = fig.add_axes([0.05, 0.65, 0.35, 0.25])
    tscore_bar(ax_lf, t_lf, "Left frontal")

    # Right frontal
    ax_rf = fig.add_axes([0.60, 0.65, 0.35, 0.25])
    tscore_bar(ax_rf, t_rf, "Right frontal")

    # Brain
    ax_brain = fig.add_axes([0.35, 0.35, 0.30, 0.30])
    ax_brain.imshow(mpimg.imread(brain_img_path))
    ax_brain.axis("off")

    # Left posterior
    ax_lp = fig.add_axes([0.05, 0.10, 0.35, 0.25])
    tscore_bar(ax_lp, t_lp, "Left posterior")

    # Right posterior
    ax_rp = fig.add_axes([0.60, 0.10, 0.35, 0.25])
    tscore_bar(ax_rp, t_rp, "Right posterior")

    plt.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close()

def plot_integrated_connectivity(t_score, brain_img_path, save_path):
    import matplotlib.pyplot as plt
    import matplotlib.image as mpimg

    fig = plt.figure(figsize=(6, 4))
    fig.patch.set_facecolor("white")

    # Brain image
    ax_brain = fig.add_axes([0.25, 0.45, 0.5, 0.45])
    ax_brain.imshow(mpimg.imread(brain_img_path))
    ax_brain.axis("off")

    # Indicator below
    ax_ind = fig.add_axes([0.10, 0.10, 0.80, 0.25])
    tscore_bar(ax_ind, t_score, "Integrative Resilience Index")

    plt.savefig(save_path, dpi=300, bbox_inches="tight")
    plt.close()


def generate_figures(results, output_directory, brain_img_path):
    plot_directory = f"{output_directory}/plots"
    figures = {}

    # ---- Integrated ----
    figures["integrated"] = os.path.join(
        plot_directory, "integrated_connectivity.png"
    )
    plot_integrated_connectivity(
        results["integrated"]["t_score"],
        brain_img_path=brain_img_path,
        save_path = figures["integrated"]
    )

    # ---- Hemispheric ----
    figures["hemispheric"] = os.path.join(
        plot_directory, "hemispheric_connectivity.png"
    )
    plot_hemispheric_composite(
        t_left=results["left_hemi"]["t_score"],
        t_right=results["right_hemi"]["t_score"],
        t_inter=results["inter_hemi"]["t_score"],
        brain_img_path=brain_img_path,
        save_path=figures["hemispheric"]
    )

    # ---- Local ----
    figures["local"] = os.path.join(
        plot_directory, "local_connectivity.png"
    )
    plot_local_composite(
        t_lf=results["left_frontal"]["t_score"],
        t_rf=results["right_frontal"]["t_score"],
        t_lp=results["left_posterior"]["t_score"],
        t_rp=results["right_posterior"]["t_score"],
        brain_img_path=brain_img_path,
        save_path=figures["local"]
    )

    return figures

def brain_connectivity(raw, subject_metadata, normative_stats, target_dir, output_path, brain_image):

    os.makedirs(target_dir, exist_ok=True)
    # -----------------------------
    # 1. PREPROCESSING
    # -----------------------------
    raw_clean = preprocess_raw(raw)

    # -----------------------------
    # 2. EPOCHING
    # -----------------------------
    epochs = make_epochs(raw_clean)
    # -----------------------------
    # 3. CONNECTIVITY (ALL BANDS)
    # -----------------------------
    band_matrices = {}
    mat = compute_band_connectivity(epochs, 8, 13)
    print("min:", np.nanmin(mat))
    print("max:", np.nanmax(mat))
    print("symmetry check:", np.nanmax(np.abs(mat - mat.T)))

    for band, (fmin, fmax) in EEG_Bands.items():
        band_matrices[band] = compute_band_connectivity(
            epochs, fmin, fmax
        )

    # -----------------------------
    # 4. REGION DEFINITIONS
    # -----------------------------
    regions = define_regions(epochs.ch_names)

    # -----------------------------
    # 5. AGGREGATE CONNECTIVITY
    # -----------------------------
    metrics = {
        "integrated": [],
        "left_hemi": [],
        "right_hemi": [],
        "inter_hemi": [],
        "left_frontal": [],
        "right_frontal": [],
        "left_posterior": [],
        "right_posterior": [],
    }

    for band, mat in band_matrices.items():
        metrics["integrated"].append(
            mean_connectivity(mat, range(mat.shape[0]))
        )
        metrics["left_hemi"].append(
            mean_connectivity(mat, regions["left_hemi"])
        )
        metrics["right_hemi"].append(
            mean_connectivity(mat, regions["right_hemi"])
        )
        metrics["inter_hemi"].append(
            mean_connectivity(
                mat,
                regions["left_hemi"],
                regions["right_hemi"]
            )
        )
        for r in ["left_frontal", "right_frontal",
                  "left_posterior", "right_posterior"]:
            metrics[r].append(
                mean_connectivity(mat, regions[r])
            )

    # Average across bands
    print('metrics', metrics)
    metrics = {k: sum(v) / len(v) for k, v in metrics.items()}
    # -----------------------------
    # 6. NORMALIZATION (T-SCORES)
    # -----------------------------
    results = {}

    for region, value in metrics.items():
        μ = normative_stats[region]["mean"]
        σ = normative_stats[region]["std"]

        t = t_score(value, μ, σ)
        print("t values", t)
        label = classify(t)

        results[region] = {
            "value": value,
            "t_score": t,
            "label": label
        }

    # -----------------------------
    # 7. FIGURE GENERATION
    # -----------------------------
#new
    figures = generate_figures(results, target_dir, brain_image)
    

    # -----------------------------
    # 8. ANALYTICS TEXT
    # -----------------------------
    
    t_scores = {
    k: v["t_score"]
    for k, v in results.items()
    }

    classes = {
        k: v["label"]
        for k, v in results.items()
    }

    sections = {
    "integrated": {
        "figure": figures["integrated"],  # integrative resilience index figure
        "t_scores": {
            "left_hemi": t_scores["left_hemi"],
            "right_hemi": t_scores["right_hemi"],
            "inter_hemi": t_scores["inter_hemi"],
            "left_frontal": t_scores["left_frontal"],
            "right_frontal": t_scores["right_frontal"],
            "left_posterior": t_scores["left_posterior"],
            "right_posterior": t_scores["right_posterior"],
        },
        "classes": {
            "left_hemi": classes["left_hemi"],
            "right_hemi": classes["right_hemi"],
            "inter_hemi": classes["inter_hemi"],
            "left_frontal": classes["left_frontal"],
            "right_frontal": classes["right_frontal"],
            "left_posterior": classes["left_posterior"],
            "right_posterior": classes["right_posterior"],
        }
    },

    "hemispheric": {
        "figure": figures["hemispheric"],  # composite hemispheric PNG
        "t_scores": {
            "left_hemi": t_scores["left_hemi"],
            "right_hemi": t_scores["right_hemi"],
            "inter_hemi": t_scores["inter_hemi"],
        },
        "classes": {
            "left_hemi": classes["left_hemi"],
            "right_hemi": classes["right_hemi"],
            "inter_hemi": classes["inter_hemi"],
        }
    },

    "local": {
        "figure": figures["local"],  # composite local PNG
        "t_scores": {   
            "left_frontal": t_scores["left_frontal"],
            "right_frontal": t_scores["right_frontal"],
            "left_posterior": t_scores["left_posterior"],
            "right_posterior": t_scores["right_posterior"],
        },
        "classes": {
            "left_frontal": classes["left_frontal"],
            "right_frontal": classes["right_frontal"],
            "left_posterior": classes["left_posterior"],
            "right_posterior": classes["right_posterior"],
            }
        }
    }


    # -----------------------------
    # 9. BUILD PDF
    # -----------------------------
    pdf_path = os.path.join(
        target_dir,
        output_path
    )

    build_pdf(pdf_path, sections, subject_metadata)
    # return False
    return pdf_path

def preprocess_raw(raw):
    if 'EXT' in raw.ch_names:
        raw.set_channel_types({'EXT': 'misc'})

    raw.filter(1., 40., fir_design="firwin")
    raw.set_eeg_reference("average", projection=True)
    raw.set_montage("standard_1020", on_missing="ignore")
    return raw

def compute_radar_features(raw):
    
    #Compute EEG features and convert them to T-scores.
    #T-score = 50 + 10 * (x - mean_norm) / std_norm
   

    psd = raw.compute_psd(fmin=1, fmax=40, picks="eeg", verbose=False)
    psds = np.log(psd.get_data())
    freqs = psd.freqs

    def band(fmin, fmax):
        return psds[:, (freqs >= fmin) & (freqs <= fmax)].mean()

    #  Raw metrics (physiological proxies)
    metrics = {}
    metrics["Stress"] = band(13, 30) / band(8, 13)
    metrics["Anxiety"] = band(20, 30) / band(8, 10)
    metrics["Attention Deficit"] = band(4, 8) / band(13, 30)

    f3 = raw.ch_names.index("F3")
    f4 = raw.ch_names.index("F4")
    alpha = (freqs >= 8) & (freqs <= 13)
    metrics["Depression"] = psds[f4, alpha].mean() - psds[f3, alpha].mean()

    mean_alpha = psds.mean(axis=0)[alpha]
    eeg_speed_hz = freqs[alpha][np.argmax(mean_alpha)]
    metrics["EEG Speed"] = eeg_speed_hz

    # Dummy normative stats (REPLACE later with real norms) 
    norm_mean = {
        "Depression": 0.0,
        "Anxiety": 1.2,
        "Stress": 1.3,
        "Attention Deficit": 1.0,
        "EEG Speed": 10.0
    }

    norm_std = {
        "Depression": 0.3,
        "Anxiety": 0.4,
        "Stress": 0.4,
        "Attention Deficit": 0.3,
        "EEG Speed": 1.0
    }

    #  Convert to T-scores
    t_scores = {}
    for k, v in metrics.items():
        t_scores[k] = np.clip(
            50 + 10 * (v - norm_mean[k]) / norm_std[k],
            20, 90
        )

    # Attach EEG speed Hz explicitly (as in PDF)
    t_scores["EEG Speed Hz"] = round(eeg_speed_hz, 1)

    return t_scores

def plot_radar(t_scores, out_png):
    labels = [
        "Depression",
        "Anxiety",
        "Stress",
        "EEG Speed",
        "Attention Deficit"
    ]

    values = [t_scores[k] for k in labels]
    values += values[:1]

    angles = np.linspace(0, 2*np.pi, len(labels), endpoint=False)
    angles = np.append(angles, angles[0])

    fig = plt.figure(figsize=(6, 6))
    ax = plt.subplot(111, polar=True)

    # Radar
    ax.plot(angles, values, linewidth=2, color="black")
    ax.fill(angles, values, alpha=0.25)

    # Axis formatting
    ax.set_thetagrids(angles[:-1]*180/np.pi, labels, fontsize=10)
    ax.set_ylim(20, 90)
    ax.set_yticks([30, 50, 70])
    ax.set_yticklabels(["Mild", "Normal", "Severe"], fontsize=8)
    ax.grid(True)

    # Value annotations
    for angle, val in zip(angles, values):
        ax.text(angle, val + 2, f"{int(val)}",
                ha="center", va="center", fontsize=9)

    plt.title("Major Brain Function Indices (T-score)", y=1.12, fontsize=12)

    plt.savefig(out_png, dpi=300, bbox_inches="tight")
    plt.close()

def generate_comments(t_scores):
    comments = {}

    def level(score):
        if score < 40:
            return "Mild"
        elif score < 60:
            return "Normal"
        elif score < 75:
            return "Moderate"
        else:
            return "Severe"

    templates = {
        "Depression": {
            "Normal": "Your brainwave patterns do not show any signs typically associated with depression.",
            "Moderate": "Your brainwave patterns show certain patterns associated with depression.",
            "Severe": "Your brainwave patterns show strong indicators associated with depression."
        },
        "Anxiety": {
            "Normal": "Your brainwave patterns do not show any signs typically associated with anxiety.",
            "Moderate": "Your brainwave patterns show certain signs associated with anxiety.",
            "Severe": "Your brainwave patterns show strong indicators associated with anxiety."
        },
        "Stress": {
            "Normal": "Your brainwave patterns do not show any signs of excessive stress.",
            "Moderate": "Your brainwave patterns show signs associated with stress.",
            "Severe": "Your brainwave patterns show strong indicators associated with stress."
        },
        "Attention Deficit": {
            "Normal": "The brainwave patterns do not show significant patterns associated with attention deficit.",
            "Moderate": "Your brainwave patterns show signs associated with attention difficulties.",
            "Severe": "Your brainwave patterns show strong indicators associated with attention deficit."
        },
        "EEG Speed": {
            "Normal": "Your brainwave speed is average, indicating normal alertness and focus.",
            "Moderate": "Your brainwave speed is slightly elevated.",
            "Severe": "Your brainwave speed is excessively high or low, indicating potential dysregulation."
        }
    }

    for k in ["Depression", "Anxiety", "Stress", "Attention Deficit", "EEG Speed"]:
        lvl = level(t_scores[k])
        comments[k] = templates[k].get(lvl, templates[k]["Normal"])

    return comments

def compute_sloreta(raw, tmin=30, tmax=40):
    
    #Memory-safe sLORETA computation on a short window.
    
     
    # Crop to representative window
    
    raw_crop = raw.copy().crop(tmin=tmin, tmax=tmax)

    # Optional decimation (HUGE memory saver)
    raw_crop.resample(100)  # from 500 Hz → 100 Hz

    
    # Forward model
    
    fwd = mne.make_forward_solution(
        raw_crop.info,
        trans="fsaverage",
        src=src,
        bem=bem_sol,
        eeg=True,
        meg=False,
        n_jobs=1
    )

  
    # Noise covariance
    
    noise_cov = mne.compute_raw_covariance(
        raw_crop,
        tmin=0,
        tmax=None,
        rank="info"
    )

   
    # Inverse operator
 
    inv = make_inverse_operator(
        raw_crop.info,
        fwd,
        noise_cov,
        loose=0.2,
        depth=0.8
    )

    # Apply inverse (SAFE)
   
    stc = apply_inverse_raw(
        raw_crop,
        inv,
        lambda2=1 / 9,
        method="sLORETA",
        buffer_size=2000   # critical memory control
    )

    return stc

def zscore_stc(stc):
    data = stc.data
    mean = np.mean(data)
    std = np.std(data)
    z_data = (data - mean) / std
    return mne.SourceEstimate(
        z_data,
        vertices=stc.vertices,
        tmin=stc.tmin,
        tstep=stc.tstep,
        subject=stc.subject
    )

def plot_axial_slices(stc_z, out_png, n_slices=10):
   
    #Generate LORETA-style axial slices from Z-scored
    #surface sLORETA using stable MNE API.
   
    # ---- Collapse time (clinical standard) ----
    stc_mean = stc_z.mean()

    # ---- Surface → volume (SAFE, low-res) ----
    vol = stc_mean.as_volume(
        src,
        subject="fsaverage",
        mri_resolution=False   # CRITICAL: prevents RAM crash
    )

    data = vol.get_fdata()

    # ---- Choose axial slices (avoid extremes) ----
    z_vals = np.linspace(
        int(data.shape[2] * 0.25),
        int(data.shape[2] * 0.75),
        n_slices
    ).astype(int)

    fig, axes = plt.subplots(2, 5, figsize=(12, 5))

    vmax = np.percentile(np.abs(data), 98)

    for ax, z in zip(axes.flat, z_vals):
        ax.imshow(
            data[:, :, z].T,
            cmap="RdBu_r",
            origin="lower",
            vmin=-vmax,
            vmax=vmax
        )
        ax.set_title(f"Axial Z={z}", fontsize=8)
        ax.axis("off")

    fig.suptitle("Axial sLORETA Z-score Maps", fontsize=12)
    plt.tight_layout(rect=[0, 0, 1, 0.95])
    plt.savefig(out_png, dpi=300)
    plt.close()

def plot_connectivity(raw, out_png, fmin=8, fmax=13):
    # Ensure output directory exists
    os.makedirs(os.path.dirname(out_png), exist_ok=True)

    # ---- Short epochs (memory-safe) ----
    epochs = mne.make_fixed_length_epochs(
        raw,
        duration=2.0,
        overlap=1.0,
        preload=True
    )

    # ---- Connectivity (NEW API) ----
    conn = spectral_connectivity_epochs(
        epochs,
        method="coh",
        mode="fourier",
        sfreq=raw.info["sfreq"],
        fmin=fmin,
        fmax=fmax,
        faverage=True,
        verbose=False
    )

    # Extract connectivity matrix
    con = conn.get_data(output="dense")[:, :, 0]  # alpha band

    # ---- Plot ----
    fig, ax = plt.subplots(figsize=(6, 6))
    im = ax.imshow(con, cmap="hot", vmin=0, vmax=1)
    ax.set_title("Alpha-band Functional Connectivity")

    ax.set_xticks(range(len(raw.ch_names)))
    ax.set_yticks(range(len(raw.ch_names)))
    ax.set_xticklabels(raw.ch_names, rotation=90, fontsize=6)
    ax.set_yticklabels(raw.ch_names, fontsize=6)

    plt.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    plt.tight_layout()
    plt.savefig(out_png, dpi=300)
    plt.close()

def extract_mni_peaks(stc, n=10):


    subjects_dir = mne.get_config("SUBJECTS_DIR")
    subject = "fsaverage"

    # Mean power across time
    power = stc.data.mean(axis=1)

    # Split hemispheres
    lh_n = len(stc.vertices[0])
    rh_n = len(stc.vertices[1])

    power_lh = power[:lh_n]
    power_rh = power[lh_n:lh_n + rh_n]

    # Get MNI coords per hemisphere
    lh_coords = mne.vertex_to_mni(
        stc.vertices[0], hemis=0, subject=subject, subjects_dir=subjects_dir
    )
    rh_coords = mne.vertex_to_mni(
        stc.vertices[1], hemis=1, subject=subject, subjects_dir=subjects_dir
    )

    # Concatenate
    coords = np.vstack([lh_coords, rh_coords])
    power_all = np.hstack([power_lh, power_rh])

    # Select top-N absolute peaks
    idx = np.argsort(np.abs(power_all))[-n:]

    return coords[idx], power_all[idx]


def generate_pdf(pdf_path, sections):
    
   # sections = list of tuples:
       # ("Title", image_path),
       # ("Analytics results", comments_dict),
        
    

    c = canvas.Canvas(pdf_path, pagesize=A4)
    W, H = A4

    def draw_wrapped_text(text, x, y, max_width, leading=14):
        lines = simpleSplit(text, "Helvetica", 10, max_width)
        for line in lines:
            c.drawString(x, y, line)
            y -= leading
        return y

    for title, content in sections:

        
        # Page Title
        
        c.setFont("Helvetica-Bold", 14)
        c.drawString(40, H - 40, title)

        
        # CASE 1: Image section
      
        if isinstance(content, str):
            c.drawImage(
                content,
                40,
                H - 520,
                width=500,
                height=400,
                preserveAspectRatio=True,
                anchor="c"
            )

        
        # CASE 2: Radar analytics text
    
        elif isinstance(content, dict):
            y = H - 90
            c.setFont("Helvetica", 10)

            for key in [
                "Depression",
                "Anxiety",
                "Stress",
                "EEG Speed",
                "Attention Deficit"
            ]:
                if key not in content:
                    continue

                # Bullet / checkmark
                c.setFillColor(HexColor("#000000"))
                c.drawString(40, y, "✓")

                # Label
                c.setFont("Helvetica-Bold", 10)
                c.drawString(55, y, key)

                # Comment text
                c.setFont("Helvetica", 10)
                y = draw_wrapped_text(
                    content[key],
                    x=130,
                    y=y,
                    max_width=420,
                    leading=14
                )

                y -= 10  # spacing between indicators

        c.showPage()

    c.save()

def plot_3d_brain(stc_z, out_png, hemi="both"):
    os.makedirs(os.path.dirname(out_png), exist_ok=True)

    brain = stc_z.plot(
        subject="fsaverage",
        subjects_dir=subjects_dir,
        hemi=hemi,
        surface="inflated",
        smoothing_steps=10,       
        time_viewer=False,
        size=(900, 700),
        background="white",
        clim=dict(
            kind="value",
            lims=[-3, 0, 3]         # Z-score scale
        )
    )

    brain.save_image(out_png)
    brain.close()

def plot_axial_slices(stc_z, out_png):

    # Ensure output directory exists
    os.makedirs(os.path.dirname(out_png), exist_ok=True)

    # Plot axial views
    brain = stc_z.plot(
        subject="fsaverage",
        subjects_dir=subjects_dir,
        hemi="both",
        views=["axial"],
        surface="inflated",
        background="white",
        colorbar=True,
        time_label="",
        size=(800, 600)
    )

    # Save screenshot
    brain.save_image(out_png)
    brain.close()

def plot_indicators(raw, basename, output_dir):

    #basename code
    parts = basename.split("_", 1)
    patient_id, patient_name = (parts[0], parts[1]) if len(parts) > 1 else (None, None)
    patient_dir = output_dir
    plots_dir = os.path.join(patient_dir, "plots")

    os.makedirs(plots_dir, exist_ok=True)


    # Load & preprocess
    

    raw = preprocess_raw(raw)
    # Radar features (T-scores)

    t_scores = compute_radar_features(raw)
    radar_comments = generate_comments(t_scores)

    radar_png = os.path.join(plots_dir, "radar.png")
    plot_radar(t_scores, radar_png)


    # Print analytical comments (console)
    
    print(f"\n=== Radar Analysis : {patient_id}:")
    for k in ["Depression", "Anxiety", "Stress", "EEG Speed", "Attention Deficit"]:
        score = int(t_scores[k])
        comment = radar_comments[k]
        print(f"{k:18s} | T-score: {score:>3} | {comment}")


    # sLORETA + Z-score
    
    stc = compute_sloreta(raw)
    stc_z = zscore_stc(stc)

    brain_png = os.path.join(plots_dir, "brain_3d.png")
    slices_png = os.path.join(plots_dir, "axial_slices.png")
    conn_png = os.path.join(plots_dir, "connectivity.png")

    plot_3d_brain(stc_z, brain_png)
    plot_axial_slices(stc_z, slices_png)
    plot_connectivity(raw, conn_png)


    # MNI peak extraction

    mni_coords, power = extract_mni_peaks(stc_z)

    
    # PDF generation (radar + comments + plots)
    
    pdf_path = os.path.join(patient_dir, f"{patient_id}_Brain_Indicators.pdf")

    generate_pdf(
        pdf_path,
        sections=[
            ("Major Brain Function Indices (T-score)", radar_png),
            ("Radar Analytics Results", radar_comments),
            ("3D Source Localization (Z-score)", brain_png),
            ("Axial Slices", slices_png),
            ("Functional Connectivity", conn_png)
        ]
    )

    print(f" Report generated: {pdf_path}\n")

def roman(num):
    vals = [
        (1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),
        (100,'C'),(90,'XC'),(50,'L'),(40,'XL'),
        (10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')
    ]
    r = ""
    for v, s in vals:
        while num >= v:
            r += s
            num -= v
    return r

#For sections without subsection

def insert_section_images(
    doc,
    image_folder,
    section_idx,
    width=Inches(7.2),
    height=Inches(3.8),
    images_per_page=2
):
    pattern = re.compile(
        rf"^{section_idx}_(\d+)_.*\.(png|jpg|jpeg)$",
        re.IGNORECASE
    )

    matched = []
    for img in os.listdir(image_folder):
        m = pattern.match(img)
        if m:
            img_num = int(m.group(1))
            matched.append((img_num, img))

    matched.sort(key=lambda x: x[0])

    for i, (_, img) in enumerate(matched):
        doc.add_picture(
            os.path.join(image_folder, img),
            width=width,
            height=height
        )

        # Page break after every 2 images
        #if (i + 1) % images_per_page == 0:
           # doc.add_page_break()

def insert_section4_images(doc, image_folder):
  
    files = sorted(os.listdir(image_folder))

    # IV-I Band
    doc.add_heading("IV-I Band", level=2)

    band_imgs = [
        f for f in files
        if re.match(r"^4_1_\d+_.*\.(png|jpg|jpeg)$", f, re.I)
    ]

    for img in band_imgs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        run.add_picture(
            os.path.join(image_folder, img),
            width=Inches(7.2),
            height=Inches(4.2)
        )

        # spacing after each image (same look as other sections)
        p.paragraph_format.space_after = Inches(0.3)
    # IV-II Absolute (full-width)
   
    doc.add_heading("IV-II Absolute band powers (1 – 45 Hz)", level=2)

    abs_imgs = [
        f for f in files
        if re.match(r"^4_2_\d+_.*\.(png|jpg|jpeg)$", f, re.I)
    ]

    for img in abs_imgs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(
            os.path.join(image_folder, img),
            width=Inches(7.2), height = Inches(4.2)
        )

   
    # IV-III Relative (full-width)
    
    doc.add_heading("IV-III Relative band powers (1 – 45 Hz)", level=2)

    rel_imgs = [
        f for f in files
        if re.match(r"^4_3_\d+_.*\.(png|jpg|jpeg)$", f, re.I)
    ]

    for img in rel_imgs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(
            os.path.join(image_folder, img),
            width=Inches(7.2), height = Inches(4.2)
        )

def add_subsection_heading(doc, main_sec, sub_sec, title):
    doc.add_heading(
        f"{roman(main_sec)}-{sub_sec} {title}",
        level=2
    )

def add_floating_image(paragraph, image_path, width, height, 
                       h_pos, h_rel, v_pos, v_rel, 
                       behind=True, allow_overlap=True, mirror = False):
    """
    paragraph: The paragraph to anchor the image to.
    h_pos: Absolute value (e.g., Cm(8.17)) OR 'center', 'left', 'right'.
    h_rel: 'page', 'margin', or 'column'.
    v_pos: Absolute value (e.g., Cm(-1.8)).
    v_rel: 'page', 'margin', 'paragraph', or 'line'.
    """
    run = paragraph.add_run()
    picture = run.add_picture(image_path, width=width, height=height)
    inline = picture._inline
    
    # Construct Horizontal XML
    if isinstance(h_pos, str):
        h_xml = f'<wp:positionH relativeFrom="{h_rel}"><wp:align>{h_pos}</wp:align></wp:positionH>'
    else:
        h_xml = f'<wp:positionH relativeFrom="{h_rel}"><wp:posOffset>{int(h_pos.emu)}</wp:posOffset></wp:positionH>'
        
    # Construct Vertical XML
    v_xml = f'<wp:positionV relativeFrom="{v_rel}"><wp:posOffset>{int(v_pos.emu)}</wp:posOffset></wp:positionV>'

    anchor_xml = (
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        f'relativeHeight="251658240" behindDoc="{1 if behind else 0}" locked="0" layoutInCell="0" allowOverlap="{1 if allow_overlap else 0}" {nsdecls("wp", "a", "pic", "r")}>'
        f'<wp:simplePos x="0" y="0"/>'
        f'{h_xml}{v_xml}'
        f'<wp:extent cx="{int(width.emu)}" cy="{int(height.emu)}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
        f'<wp:docPr id="1" name="Picture"/><wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        f'</wp:anchor>'
    )
    
    anchor = parse_xml(anchor_xml)
    anchor.append(inline.graphic)
    if mirror:
        try:
            xfrm = inline.graphic.xpath('.//a:xfrm', namespaces={'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})[0]
            xfrm.set('flipH', '1') # This mirrors the image horizontally
        except IndexError:
            pass # Should not happen if image added correctly
    inline.getparent().replace(inline, anchor)

def make_doc(base_name, raw, image_folder_path, output_doc, icons_path):
    

    plot_image_path = Path(f"{image_folder_path}/plots")
    plot_image_path.mkdir(exist_ok=True)
    # print(plot_image_path)
    icons_path = Path(f"{icons_path}")
    icons_path.mkdir(exist_ok=True)
    print(icons_path)
    parts = base_name.split("_", 1)
    patient_id, patient_name = (parts[0], parts[1]) if len(parts) > 1 else (None, None)

    measurement_date = raw.info.get("meas_date")
    n_channels = raw.info["nchan"]
    channel_names = ", ".join(raw.info["ch_names"])
    sfreq = raw.info["sfreq"]
    

    metadata = {"name": patient_name,
                "id": patient_id,
                "meas_date": measurement_date,
                "n_ch": n_channels,
                "ch_names": channel_names,
                "sfreq": sfreq,
                "sex": "Male",
                "date_of_birth": "test",
                "age": 41
                }

    styles = getSampleStyleSheet()

    doc = Document()
    style_title = doc.styles['Title']
    font = style_title.font
    font.color.rgb = RGBColor(0, 147, 214)
    font.name = 'PF Agora Sans Pro'

    style_heading1 = doc.styles['Heading 1']
    font = style_heading1.font
    font.color.rgb = RGBColor(0, 147, 214)
    font.name = 'PF Agora Sans Pro'

    style_heading2 = doc.styles['Heading 2']
    font = style_heading2.font
    font.color.rgb = RGBColor(227, 97, 0)
    font.name = 'PF Agora Sans Pro'

    section = doc.sections[0]

       # Margins (tight but safe for PDF)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Header/Footer distance
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.5)


    # first page header
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False

    table = first_header.add_table(rows=1, cols=3, width=Inches(9))
    table.autofit = False
    left_cell = table.cell(0, 0)
    center_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    left_cell.width = Inches(2.33)
    center_cell.width = Inches(2.8)
    right_cell.width = Inches(2.8)


    p_center = center_cell.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p_center.add_run("neurowellness")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 147, 214)
    run.font.name = "Montserrat"

    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # run.add_break(WD_BREAK.LINE)
    # run = p_right.add_run()
    # run.text = measurement_date.strftime("%d-%m-%Y")
    run = p_right.add_run()
    add_floating_image(
    p_center, 
    rf"{icons_path}/Lines5.png", 
    width=Cm(6.45), 
    height=Cm(6.37),
    h_pos=Cm(13.82),
    h_rel = "column",
    v_pos = Cm(-1.2),
    v_rel = "paragraph"
)   
    # Main header
    header = section.header
    header.is_linked_to_previous = False

    table = header.add_table(rows=1, cols=3, width=Inches(7))
    table.autofit = False


    left_cell = table.cell(0, 0)
    center_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    left_cell.width = Inches(2.33)
    center_cell.width = Inches(2.8)
    right_cell.width = Inches(2.33)

    p_left = left_cell.paragraphs[0]
    p_left.text = patient_name
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p_right = right_cell.paragraphs[0]
    # p_right.text = measurement_date.strftime("%d-%m-%Y")
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_center = center_cell.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p_center.add_run("neurowellness")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0, 147, 214)
    run.font.name = "Montserrat"
    run.add_break(WD_BREAK.LINE)
    run = p_right.add_run()
    run.text = measurement_date.strftime("%d-%m-%Y")

  

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    # First footer
    first_footer = section.first_page_footer

    table = first_footer.add_table(rows=1, cols=3, width=Inches(7.0)) # Adjusted slightly for standard margins
    table.autofit = False

    # Define the cells
    left_cell = table.cell(0, 0)
    center_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    # Set widths (Dividing roughly 7 inches by 3)
    left_cell.width = Inches(2.8)
    center_cell.width = Inches(2.8)
    right_cell.width = Inches(2.33)

    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p_left.add_run()
    add_floating_image(
    p_left,
    rf"{icons_path}/Graphics 1.png",
    width=Cm(7.93), 
    height=Cm(3.31),
    h_pos=Cm(-2.14),
    h_rel = "page",
    v_pos = Cm(-1.52),
    v_rel = "paragraph",

    )
    # --- YOUR CENTERED TEXT SECTION ---
    p_center = center_cell.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Line 1: Neuromodulation
    run1 = p_center.add_run("Neuromodulation ")
    run1.bold = True
    run1.font.size = Pt(9)
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.name = "Montserrat"   # or "Arial Black" if Montserrat not installed


    # Line 2: tailored for you.
    run2 = p_center.add_run("tailored for you.")
    run2.bold = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(30, 150, 190)  # blue tone
    run2.font.name = "Montserrat"


    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Page X of Y
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.extend([fldChar1, instrText, fldChar2])
    p.add_run(" of ")

    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.extend([fldChar1, instrText, fldChar2])

    # Main footer
    footer = section.footer
    footer.is_linked_to_previous = False

    table = footer.add_table(rows=1, cols=3, width=Inches(7.0)) # Adjusted slightly for standard margins
    table.autofit = False

    # Define the cells
    left_cell = table.cell(0, 0)
    center_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    # Set widths (Dividing roughly 7 inches by 3)
    left_cell.width = Inches(2.33)
    center_cell.width = Inches(2.8)
    right_cell.width = Inches(2.33)
    p_left = left_cell.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p_left.add_run()
    add_floating_image(
    p_left,
    rf"{icons_path}/Graphics 1.png",
    width=Cm(7.93), 
    height=Cm(3.31),
    h_pos=Cm(-2.14),
    h_rel = "page",
    v_pos = Cm(-1.52),
    v_rel = "paragraph",

    )
    # --- YOUR CENTERED TEXT SECTION ---
    p_center = center_cell.paragraphs[0]
    p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Line 1: Neuromodulation
    run1 = p_center.add_run("Neuromodulation ")
    run1.bold = True
    run1.font.size = Pt(9)
    run1.font.color.rgb = RGBColor(0, 0, 0)
    run1.font.name = "Montserrat"   # or "Arial Black" if Montserrat not installed

    # Line 2: tailored for you.
    run2 = p_center.add_run("tailored for you.")
    run2.bold = True
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(30, 150, 190)  # blue tone
    run2.font.name = "Montserrat"


    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Page X of Y
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.extend([fldChar1, instrText, fldChar2])
    p.add_run(" of ")

    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.extend([fldChar1, instrText, fldChar2])


    # Main Doc
    title = doc.add_heading("EEG Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Inches(0.2)

    sec_idx = 1

    # SECTION I 
    if sec_idx == 1:
        doc.add_heading(f"{roman(sec_idx)}. Patient Data", level=1)
        doc.add_paragraph(f"Patient Name : {patient_name}")
        doc.add_paragraph(f"Patient ID : {patient_id}")
        doc.add_paragraph(f"Measurement Date : {measurement_date}")
        doc.add_paragraph(f"Sampling Frequency : {sfreq} Hz")
        doc.add_paragraph(f"Number of Channels : {n_channels}")
        doc.add_paragraph(f"Channel Names : {channel_names}")

    doc.add_page_break()
    # SECTION II 
    sec_idx += 1
    if sec_idx == 2:
        doc.add_heading(f"{roman(sec_idx)}. EEG", level=1)
        doc.add_paragraph(
            "EEG data is a record of the oscillations of electrical brain potentials recorded from electrodes on the human scalp (T100) (T101) (T005).The raw data in the figure below have been cleaned by the application of high-pass and low-pass filters."
        )

        # INSERT IMAGES HERE 
        insert_section_images(
            doc=doc,
            image_folder=plot_image_path,
            section_idx=sec_idx )
        
    sec_idx+=1
    if sec_idx ==3:
        doc.add_heading(f"{roman(sec_idx)}. ICA Component Analysis", level=1)
        doc.add_paragraph(
            "Independent component analysis (ICA) is a statistical method to separate independent sources from superimposed signals. It is the most common method that has been used in EEG data decomposition, and can be used to identify and remove the artifacts from raw EEG data. Features including time series, power spectrum density (PSD), component scalp map (Topomap), dipole source location (Source location) extracted from ICA are shown for each component."
        )
        insert_section_images(
            doc=doc,
            image_folder=plot_image_path,
            section_idx=sec_idx )

    sec_idx+=1
    if sec_idx ==4:
        doc.add_heading(f"{roman(sec_idx)}. Band power - Topomap", level=1)
        doc.add_paragraph(
            "In power spectral density (PSD) 2D map, topomaps of absolute and relative power in 1 Hz bins (1 - 45 Hz) as well as each frequency band are presented. Absolute power is the sum of the component powers for each frequency band. Relative power is the absolute power in a specific frequency band divided by the total power. It is advisable to compare relative power with absolute power, since absolute power reflects the individual differences due to variations in brain tissue. This feature provides absolute and relative power based on six brain regions (prefrontal, frontal, left temporal, right temporal, central, parietal, and occipital). The power spectra for each of the 19 channels are shown in the following feature, PSD spectrum (below).")

        insert_section4_images(doc, image_folder=plot_image_path)
    doc_dir = f"{image_folder_path}/{output_doc}"
    doc.save(doc_dir)

    return metadata

def create_workspace_directory(dir_name: str):
    """
    Creates a directory in the current working directory.
    """
    try:
        Path(dir_name).mkdir(parents=True, exist_ok=True)
        print(f"Successfully ensured directory '{dir_name}' exists.")
    except Exception as e:
        print(f"An error occurred while creating the directory: {e}")

def process_nedf_files(extension: str, target_dir: str):
    """
    Searches for files with a specific extension in the current workspace.
    Example extension: '.txt' or '.pdf'
    """
    # Ensure the extension starts with a dot
    ext = extension if extension.startswith('.') else f".{extension}"
    
    # Get a list of all files in the current directory matching the extension
    files = list(CURRENT_DIR.glob(f"*{ext}"))
    
    if files:
        print(f"Found {len(files)} file(s) with extension '{ext}':")
        for file in files:
            base_name = os.path.splitext(os.path.basename(file))[0]
            parts = base_name.split("_", 1)
            patient_id, patient_name = (parts[0], parts[1]) if len(parts) > 1 else (None, None)
            target_dir = patient_name + REPORT_DIR
            print(target_dir)
            create_workspace_directory(target_dir)            
            raw = mne.io.read_raw_nedf(file, preload = True)
            montage = raw.set_montage("standard_1020", on_missing="ignore")
        
            df = raw.to_data_frame()
            
            all_dfs[base_name] = df
            raw_mne_objs[base_name] = raw
            
            
        return (all_dfs, raw_mne_objs, target_dir)
    else:
        print(f"No files with extension '{ext}' were found in the current workspace.")

def save_raw_and_cleaned_data(raw, target_dir: str = "my_dir" ):
    plots_path = 'plots'
    workspace = Path(".")
    destination = Path(f"{target_dir}/{plots_path}")
    destination.mkdir(exist_ok=True)
    raw_path = '2_1_EEG_raw.png'
    target_raw_path = destination / raw_path
    bandpass_path = '2_2_EEG_bandpass.png'
    target_bandpass_path = destination / bandpass_path
    clean_path = '2_3_EEG_clean_data.png'
    target_clean_path = destination / clean_path

    
    raw_plot = raw.copy().pick("eeg").plot(
    n_channels=20,
    duration=15,   # seconds
    start = 45,
    scalings="45e-6",
    title="Raw EEG (All Channels)",
    show=False,
    clipping = 2,
    show_scrollbars = False
    )
    raw_plot.suptitle("Raw data", y=1, x=0.5, fontsize=16)
    raw_plot.savefig(target_raw_path, dpi=300)
    plt.close()
    raw_filtered = raw.filter(l_freq = 0.5, h_freq = 45.0, fir_design = 'firwin')
    raw_badpass = raw_filtered.copy().pick("eeg").plot(
    n_channels=20,
    duration=15,   # seconds
    start = 45,
    scalings="45e-6",
    title="Raw EEG (All Channels)",
    show=False,
    clipping = 2,
    show_scrollbars = False
    )

    raw_badpass.suptitle("Raw data (Bandpass filered)", y=1, x=0.5, fontsize=16)
    raw_badpass.savefig(target_bandpass_path, dpi=300)
    plt.close()

    raw_average = raw_filtered.copy().set_eeg_reference("average", projection=False)

    raw_filt = raw_average.filter(
            l_freq=0.5,
            h_freq=45.0,                                      #Band Pass filtering
            picks="eeg",
            verbose=False
        )

    raw_average_plot = raw_filt.plot(
        n_channels=20,
        duration=15,
        title="Filtered EEG (0.5 - 45 Hz)",
        show=False,
        scalings = "45e-6",
        start =45,
        color = 'darkblue',
        show_scrollbars = False
    )
    raw_average_plot.suptitle("Cleaned data (Common average reference)", y=1, x=0.5, fontsize=16)
    
    raw_average_plot.savefig(target_clean_path, dpi=300)
    plt.close()

    return raw_filt

def plot_component_analysis(raw, ica, component_idx,  dipole_list, trans, destination, tmin=0, tmax=10, cmap = 'viridis'):
    """
    Plots a complete analysis figure for a single ICA component.
    """
    # Create the figure with a complex layout
    # fig = plt.figure(figsize=(18, 9), constrained_layout=True)
    fig = plt.figure(figsize=(18, 9), constrained_layout=True)
    spec = fig.add_gridspec(3, 6, height_ratios=[1, 2.5, 2.5]) # 2 rows, 6 columns for better spacing

    # --- 1. TIME SERIES (Top Row, spanning full width) ---
    ax_time = fig.add_subplot(spec[0, :])
    # Crop to 10s window
    # raw_segment = raw.copy().crop(tmin=tmin, tmax=tmax)
    # times = raw_segment.times
    raw_10s = raw.copy().crop(tmin=tmin, tmax=tmax)
    reconstructed = ica.apply(raw_10s.copy(), include=[component_idx], verbose=False)
    
    # We pick the channel where this component is strongest to show the scale
    # Alternatively, you can pick a fixed channel like 'Fz' or 'Oz'
    data_uv = reconstructed.get_data()[0] # Convert Volts to uV
    times = reconstructed.times
    
    ax_time.plot(times, data_uv, color='black', lw=0.6)
    ax_time.set_title(f'Time series - [ Component {component_idx + 1} ]', loc='center', fontweight='bold')
    ax_time.set_ylabel(r'EEG ($\mu V$)')
    ax_time.set_xlabel('Time (s)')
    # ax_time.set_ylim([-45, 45])
    # ax_time.set_yticks(np.arange(-45, 45+ 5, 40))
    ax_time.set_xlim([tmin, tmax])
    ax_time.set_xticks(np.arange(tmin, tmax + 1, 2))
    ax_time.grid(True, axis='x', color='lightgrey', lw=0.5)

    # --- 2. POWER SPECTRUM (Bottom Row, Left) ---
    ax_psd = fig.add_subplot(spec[1, 0:2])
    # Compute PSD for the specific IC source
    sources = ica.get_sources(raw)
    psd = sources.compute_psd(fmin=1, fmax=45, method='welch', picks = 'all')
    psd_data = psd.get_data()[component_idx]
    
    ax_psd.plot(psd.freqs, psd_data, color='black', lw=1)
    ax_psd.set_title('Power spectrum')
    ax_psd.set_xlabel('Frequency (Hz)')
    ax_psd.set_ylabel(r'Power ($\mu V^2/Hz$)')
    ax_psd.grid(True, alpha=0.3)

    # --- 3. TOPOMAP (Bottom Row, Center-Left) ---
    ax_topo = fig.add_subplot(spec[1, 2])
    # Get the scalp field distribution
    ica_map = ica.get_components()[:, component_idx]
    plot_topomap(ica_map, raw.info, axes=ax_topo, show=False, cmap=cmap)
    ax_topo.set_title('Topomap')

    # --- 4. DIPOLE ORTHOVIEW (Bottom Row, Right) ---
    # We pick the dipole corresponding to this component index
    print('component id', component_idx)

    current_dipole = dipole_list[component_idx]
    
    # Create subplots for Orthoview (Sagittal, Coronal, Axial)
    # Note: MNE's plot_locations can be used, but for exact 'Orthoview' 
    # as in your image, we typically use the MRI crop or 'mne_gui'
    # Here we plot the 3D location onto a standard brain template
    # dipole_spec = spec[1, 3:].subgridspec(1, 3)
    # ax_dip_list = [fig.add_subplot(dipole_spec[0, j]) for j in range(3)]
    # current_dipole.plot_locations(trans=trans, subject='fsaverage', mode='outlines', ax = ax_dip_list)


    aseg_path = fsaverage_path / "mri" / "aparc+aseg.mgz"

    aseg_img = nib.load(aseg_path)
    aseg_data = aseg_img.get_fdata()

    dip_mni_mm = current_dipole.pos[0] * 1000  # m → mm
    x_mni, y_mni, z_mni = dip_mni_mm
    

    ax_mni = fig.add_subplot(spec[1, 5])
    ax_mni.axis("off")
    mni_text = (
    "MNI information\n\n"
    "MNI coordinates\n"
    f"X = {x_mni:.0f}\n"
    f"Y = {y_mni:.0f}\n"
    f"Z = {z_mni:.0f}\n\n"
    # f"Distance = {distance_mm:.2f} mm\n"
    # f"{brodmann}\n"
    # f"{anat_label.replace('ctx-', '').replace('-', ' ').title()}"
    )
    ax_mni.text(
    0.5, 0.5,
    mni_text,
    transform=ax_mni.transAxes,   # 🔥 CRITICAL
    va="center",
    ha="center",
    fontsize=11,
    zorder=10,                    # 🔥 ensure on top
    bbox=dict(
        boxstyle="round,pad=0.5",
        edgecolor="red",
        linewidth=1.5,
        facecolor="whitesmoke"
    )
)  
    
    dipole_spec_outlines = spec[2, 0:3].subgridspec(1, 3)
    ax_dip_list = [fig.add_subplot(dipole_spec_outlines[0, j]) for j in range(3)]
    current_dipole.plot_locations(trans=trans, subject='fsaverage', mode='outlines', ax = ax_dip_list)
    for ax, label in zip(ax_dip_list, ['Sagittal', 'Coronal', 'Axial']):
       ax.set_title(label, fontsize=10)
    
    dipole_spec_3d = spec[1, 3:5].subgridspec(1, 1)
    ax3d = fig.add_subplot(dipole_spec_3d[0, 0], projection="3d")
    current_dipole.plot_locations(trans=trans, subject='fsaverage', mode='orthoview', ax = ax3d)

    # Dipole position in meters → mm
  

    ica_path = f'3_{component_idx+1}_IC_{component_idx:03d}_analysis.png'
    target_ica_path = destination / ica_path

    fig.savefig(target_ica_path, dpi=300)
    plt.close() # Close the plot to save memory during the loop

def calc_dipoles(ica, components, ica_eeg_info, cov, bem, trans):
    picks = mne.pick_types(ica.info, eeg=True, exclude='bads')
    for i in range(ica.n_components_):
    # for i in range(1):
        print(i)
        comp_data = components[:, [i]]
        
        comp_evoked = mne.EvokedArray(comp_data, ica_eeg_info, tmin=0)
        dip, residual = mne.fit_dipole(comp_evoked, cov, bem, trans)
        dipoles.append(dip)

        print(f"Component {i}: GOF ( Goodness of Fit) = {dip.gof[0]:.2f}%")
    return dipoles
        
def save_ica_components(raw, dipoles, target_dir: str = 'my_dir'):
    plots_path = "plots"
    destination = Path(f"{target_dir}/{plots_path}")
    destination.mkdir(exist_ok=True)
        
    raw_good_channels = raw.copy()
    if 'EXT' in raw_good_channels.ch_names:
        raw_good_channels.set_channel_types({'EXT': 'misc'})

    raw_good_channels.pick_types(eeg=True, exclude='bads')

    # match_case=False: maps EDF mixed-case names (Af7, Fc3, Cp1) to the
    # montage equivalents (AF7, FC3, CP1) that standard_1020 actually uses.
    montage = mne.channels.make_standard_montage('standard_1020')
    raw_good_channels.set_montage(montage, match_case=False, on_missing="ignore")

    # Drop channels whose 3-D position is zero OR NaN after montage assignment.
    # np.any(NaN) == True so a NaN-loc channel would pass a plain `not np.any`
    # check — catch both explicitly.
    def _has_valid_loc(ch):
        loc = ch['loc'][:3]
        return np.any(loc) and not np.any(np.isnan(loc))

    no_loc = [ch['ch_name'] for ch in raw_good_channels.info['chs']
              if not _has_valid_loc(ch)]
    if no_loc:
        print(f"ICA: dropping {len(no_loc)} channels with missing positions: {no_loc}")
        raw_good_channels.drop_channels(no_loc)

    n_components = min(19, len(raw_good_channels.ch_names) - 1)
    random_state = 42

    ica = ICA(
        n_components=n_components,
        method="fastica",
        random_state=random_state,
        max_iter="auto",
    )
    ica.fit(raw_good_channels)

    # Build ica_eeg_info directly from raw_good_channels — guaranteed valid locs.
    eeg_picks = mne.pick_types(raw_good_channels.info, eeg=True, exclude='bads')
    ica_eeg_info = mne.pick_info(raw_good_channels.info, eeg_picks)
    with ica_eeg_info._unlock():
        ica_eeg_info['dev_head_t'] = mne.transforms.Transform('meg', 'head', np.eye(4))

    # components rows correspond 1-to-1 with ica_eeg_info channels.
    # Run a final sync: if any channel slipped through with an invalid loc,
    # drop it from both components and ica_eeg_info so EvokedArray never
    # sees a missing location.
    components = ica.get_components()   # shape: (n_channels, n_components_)
    valid_idx = [j for j, ch in enumerate(ica_eeg_info['chs'])
                 if _has_valid_loc(ch)]
    if len(valid_idx) < len(ica_eeg_info['chs']):
        dropped = [ica_eeg_info['chs'][j]['ch_name']
                   for j in range(len(ica_eeg_info['chs'])) if j not in valid_idx]
        print(f"Dipole fit: removing {len(dropped)} channels still missing locs: {dropped}")
        ica_eeg_info = mne.pick_info(ica_eeg_info, valid_idx)
        components = components[valid_idx, :]

    # Covariance built from the same info as ica_eeg_info — channel lists match.
    cov = mne.make_ad_hoc_cov(ica_eeg_info)

    components = ica.get_components()
    
    # getting the dipoles for source location

    if not len(dipoles):
        dipoles = calc_dipoles(ica, components, ica_eeg_info, cov, bem_sol, trans)
    else:
        dipoles = dipoles
   

    
    print(f"Generating reports for {n_components} components...")

    for i in range(n_components):
        # We call your existing plotting function
        # (Make sure the function returns the 'fig' object)
        plot_component_analysis(raw_good_channels, ica, i, dipoles, trans, destination)
        
        
        
        if (i + 1) % 5 == 0:
            print(f"Progress: {i + 1}/{n_components} components completed.")

    print(f"All reports saved to: {os.path.abspath(target_dir)}")

def plot_band_topomaps(raw, bands, psd_data, freqs, normative_stats, target_dir = 'my_dir', is_relative: bool = False, total_power = None, cmap = "coolwarm"):
    import math 
    n_bands = len(bands)
    n_cols = 4
    n_rows = math.ceil(n_bands / n_cols)
    power_name = "Relative" if is_relative else "Absolute"
    plots = "plots"
    destination = Path(f"{target_dir}/{plots}")
    destination.mkdir(exist_ok=True)
    naming_index= "2" if is_relative else "1"
    power_path = f'4_1_{naming_index}_Bands_{power_name}.png'
    target_power_path = destination / power_path

    fig, axes = plt.subplots(
        n_rows,
        n_cols,
        figsize=(3* n_cols, 3 * n_rows),
        constrained_layout=True
    )

    axes = axes.flatten()
    # vlim=(0, total_power.max() * 0.2)
    # -----------------------
    # Band loop
    # -----------------------
    try:
        from scipy.integrate import simpson as simps
    except ImportError:
        from scipy.integrate import simps

    for i, (band_name, (fmin, fmax)) in enumerate(bands.items()):
        # 1. Mask the frequencies for the specific band
        idx = (freqs >= fmin) & (freqs <= fmax)
        freq_band = freqs[idx]
        psd_band = psd_data[:, idx] # All channels for this band

        # 2. Integrate to find total power in the band (Linear scale)
        # We do this per channel
        band_power_linear = simps(psd_band, freq_band, axis=-1)

        # 3. Convert to dB (10 * log10)
        # This is the step that matches your normative -110.78 dB values
        patient_db = 10 * np.log10(np.maximum(band_power_linear, 1e-12))
        # 4. Calculate Z-Score
        # This compares the patient's dB value to the healthy dB mean
        mean_norm = normative_stats[band_name]["mean"]
        std_norm = normative_stats[band_name]["std"]
        
        z_scores = (patient_db - mean_norm) / std_norm
        if is_relative:
            rel_power = band_power_linear / total_power
            z_scores = (rel_power - mean_norm) / std_norm

        mne.viz.plot_topomap(
            z_scores,
            raw.info,
            axes=axes[i],
            show=False,
            # contours=6,
            cmap=cmap,
            # vlim =vlim
        )

        axes[i].set_title(band_name, fontsize=11)
    fig.suptitle(f"Band Topomaps ({power_name} Power)", fontsize=16)
    for ax in axes[n_bands:]:
        ax.axis("off")
    
    fig.savefig(target_power_path, dpi=300)

    plt.close()

def plot_frequency_maps(raw, psd_data, freqs, normative_stats_freq, target_dir = 'my_dir', is_relative: bool = False, total_power = None, cmap = "coolwarm"):
    import math 
    n_tops = 45
    n_cols = 9
    n_rows = math.ceil(n_tops / n_cols)
    plots = "plots"
    power_name = "Relative" if is_relative else "Absolute"
    destination = Path(f"{target_dir}/{plots}")
    destination.mkdir(exist_ok=True)
    naming_index= "3" if is_relative else "2"
    power_path = f'4_{naming_index}_1_{power_name} Power (1-45).png'
    target_power_path = destination / power_path

    fig, axes = plt.subplots(n_rows, n_cols, figsize=(18, 10))
    axes = axes.flatten()
    eps = 1e-20
    for i, f in enumerate(range(1, 46)):
        idx = np.argmin(np.abs(freqs - f))
        "---------"
        # band_power_linear = simps(psd_band, freq_band, axis=-1)
        band_power = psd_data[:, idx]

        # 3. Convert to dB (10 * log10)
        # This is the step that matches your normative -110.78 dB values
        patient_db = 10 * np.log10(np.maximum(band_power, 1e-12))
        # 4. Calculate Z-Score
        # This compares the patient's dB value to the healthy dB mean
        mean_norm = normative_stats_freq[f]["mean"]
        std_norm = normative_stats_freq[f]["std"]
        
        z_scores = (patient_db - mean_norm) / std_norm
        if is_relative:
            rel_power = band_power / total_power
            z_scores = (rel_power - mean_norm) / std_norm
        "---------"
           # raw PSD at frequency
            
        mne.viz.plot_topomap(
            z_scores,
            raw.info,
            axes=axes[i],
            show=False,
            cmap="coolwarm",
            contours=0
        )
 
        axes[i].set_title(f"{f} Hz", fontsize=10)
    fig.suptitle(f"{power_name} Power (1 - 45hz)", fontsize=16)
    for ax in axes[n_tops:]:
        ax.axis("off")
    
    fig.savefig(target_power_path, dpi=300)

    plt.close()

def band_topomaps(
    raw,
    target_dir = 'my_dir',
    bands=None,
    fmin=1,
    fmax=250,
    n_fft=2048,
    cmap="coolwarm"
):
    """
    Plot relative band-power topomaps for all bands in ONE figure.

    Parameters
    ----------
    raw : mne.io.Raw
        Preloaded Raw EEG object
    bands : dict or None
        Frequency bands as {"Name": (fmin, fmax)}
    fmin, fmax : float
        Frequency range for total power
    n_fft : int
        FFT length for PSD
    cmap : str
        Matplotlib colormap
    """

    # -----------------------
    # Defaults
    # -----------------------
    
    
    if bands is None:
        bands = {
            "Delta (1–4 Hz)": (1, 4),
            "Theta (4–8 Hz)": (4, 8),
            "Alpha (8–12 Hz)": (8, 12),
            "Beta (13–30 Hz)": (13, 30),
            "Gamma (30–45 Hz)": (30, 45),
        }

    # -----------------------
    # Preprocessing
    # -----------------------
    raw = raw.copy().load_data()
    raw.pick_types(eeg=True)

    if raw.get_montage() is None:
        raise RuntimeError("Montage must be set for topomaps.")

    raw.set_eeg_reference("average")

    # -----------------------
    # PSD (once!)
    # -----------------------
    psd = raw.compute_psd(
        method="welch",
        fmin=fmin,
        fmax=fmax,
        n_fft=n_fft
    )
    psd_data = psd.get_data()    # (n_channels, n_freqs)
    freqs = psd.freqs

   
    try:
        from scipy.integrate import simpson as simps
    except ImportError:
        from scipy.integrate import simps
        
    # total_power = np.trapezoid(psd_data, freqs, axis=1)
    total_idx = (freqs >= 1) & (freqs <= 40)
    total_power = simps(psd_data[:, total_idx], freqs[total_idx], axis=-1)
    # print(total_power)
    



    plot_band_topomaps(raw, bands, psd_data, freqs, normative_abs_log, target_dir = target_dir, cmap=cmap)
    plot_band_topomaps(raw, bands, psd_data, freqs, normative_rel, target_dir= target_dir, is_relative = True,
                        total_power=total_power, cmap=cmap)

    plot_frequency_maps(raw, psd_data, freqs, normative_abs_log_freq, target_dir = target_dir, cmap=cmap)
    plot_frequency_maps(raw, psd_data, freqs, normative_rel_freq, target_dir= target_dir, is_relative = True,
                        total_power=total_power, cmap=cmap)

                        
def main():
    # Define your directory name here
    
    
    # Call the function
    
    _, raw_mne_objs, target_dir = process_nedf_files(".nedf", TARGET_DIR)
    for base_name, raw in raw_mne_objs.items():
        raw_clean = save_raw_and_cleaned_data(raw, target_dir)
        save_ica_components(raw_clean, dipoles, target_dir)
        band_topomaps(raw_clean, target_dir=target_dir, bands = EEG_Bands )
        metadata = make_doc(base_name, raw, target_dir, OUTPUT_DOCX, ICONS_DIR)
        doc_to_pdf(OUTPUT_DOCX, OUTPUT_PDF, target_dir)
        plot_indicators(raw, base_name, target_dir)
        brain_connectivity(raw, metadata, normative_stats, target_dir, OUTPUT_BRAIN_CONNECTIVITY, brain_image)
if __name__ == "__main__":
    main()